"""
Document Processing & Embedding Pipeline — Multimodal Edition
--------------------------------------------------------------
Supports text, images, and tables extracted from PDF, DOCX, and TXT files.

Modality handling:
  - Text        → chunked with RecursiveCharacterTextSplitter as before
  - Images      → extracted per page, sent to GPT-4.1 mini (vision) for captioning,
                  caption stored as a chunk with modality metadata
  - Tables      → extracted per page (pdfplumber for PDF, python-docx for DOCX),
                  converted to Markdown, stored as a chunk with modality metadata
  - Scanned PDF → pages with no text layer are rasterized to PNG and sent to
                  GPT-4.1 mini with an OCR-focused prompt; result chunked as modality="ocr"
                  Partial scans are handled per-page (mixed docs work correctly)

Key design decisions:
  - GPT-4.1 mini is used for image captioning (supports image input per OpenAI docs)
  - All chunks — regardless of modality — go through the same OpenAI text embedding
    and Pinecone upsert path, keeping the retrieval side unchanged
  - Metadata order preserved: document_id → chat_id → user_id (+ modality, page)
  - Max upload size raised to 12 MB
  - All other improvements from the original (retry, batching, guards, etc.) retained
"""

import asyncio
import base64
import io
import json
import logging
import time
from datetime import datetime, timezone
from functools import wraps
from typing import List, Optional

import pdfplumber
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from langchain_core.documents import Document
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from langchain_text_splitters import RecursiveCharacterTextSplitter
from PIL import Image
import re
from app.core.vector_store import vector_store
from app.models.document import DocumentModel

# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------

CHUNK_SIZE         = 500
CHUNK_OVERLAP      = 100
EMBED_BATCH_SIZE   = 100
MAX_CHUNKS_GUARD   = 10_000
MAX_FILE_SIZE_MB   = 12
MAX_FILE_SIZE_BYTES = MAX_FILE_SIZE_MB * 1024 * 1024
RETRY_ATTEMPTS     = 3
RETRY_BACKOFF      = 2.0

SUPPORTED_TYPES    = {"pdf", "docx", "txt"}

# Minimum image dimensions — skip tiny icons/decorations
MIN_IMAGE_WIDTH    = 50
MIN_IMAGE_HEIGHT   = 50

# DPI used when rasterizing scanned PDF pages for OCR
# 150 DPI: good balance between GPT-4.1 mini readability and payload size
OCR_RASTER_DPI     = 150

# DPI for full-page screenshots sent for image analysis
# Higher than OCR — diagrams need more detail to be legible
PAGE_RENDER_DPI    = 150

# ---------------------------------------------------------------------------
# Module-level singletons
# ---------------------------------------------------------------------------

logger = logging.getLogger(__name__)

_embeddings = OpenAIEmbeddings()

_splitter = RecursiveCharacterTextSplitter(
    chunk_size=CHUNK_SIZE,
    chunk_overlap=CHUNK_OVERLAP,
)

# GPT-4.1 mini — used for image analysis and OCR (supports image input)
# max_tokens=2048 gives enough room for detailed multi-figure descriptions
_vision_llm = ChatOpenAI(model="gpt-4.1-mini", max_tokens=2048)

# ---------------------------------------------------------------------------
# Helpers — general
# ---------------------------------------------------------------------------


def _utcnow() -> datetime:
    return datetime.now(timezone.utc)


def _with_retry(attempts: int = RETRY_ATTEMPTS, backoff: float = RETRY_BACKOFF):
    def decorator(fn):
        @wraps(fn)
        async def wrapper(*args, **kwargs):
            last_exc = None
            for attempt in range(1, attempts + 1):
                try:
                    return await fn(*args, **kwargs)
                except Exception as exc:
                    last_exc = exc
                    if attempt < attempts:
                        wait = backoff ** (attempt - 1)
                        logger.warning(
                            "Attempt %d/%d for %s failed — retrying in %.1fs. Error: %s",
                            attempt, attempts, fn.__name__, wait, exc,
                        )
                        await asyncio.sleep(wait)
            raise last_exc
        return wrapper
    return decorator


def _filter_chunks(chunks: List[Document]) -> List[Document]:
    filtered = [c for c in chunks if c.page_content.strip()]
    dropped = len(chunks) - len(filtered)
    if dropped:
        logger.info("Dropped %d empty/whitespace chunk(s).", dropped)
    return filtered


def _inject_metadata(chunks: List[Document], document) -> List[Document]:
    """
    Merge document-level metadata into every chunk.
    Order is preserved: document_id → chat_id → user_id
    Additional modality fields (modality, page) are appended after.
    """
    for chunk in chunks:
        extra = {
            "document_id": str(document.id),
            "chat_id":     str(document.chat_id),
            "user_id":     str(document.user_id),
        }
        # Preserve any modality/page fields already on the chunk, appended after core keys
        modality_fields = {
            k: chunk.metadata[k]
            for k in ("modality", "page")
            if k in chunk.metadata
        }
        chunk.metadata = {**extra, **modality_fields}
    return chunks


# ---------------------------------------------------------------------------
# Full-page figure analysis via GPT-4.1 mini (vision)
# ---------------------------------------------------------------------------

_FIGURE_ANALYSIS_SYSTEM_PROMPT = """\
You are a document analysis assistant processing pages from an academic or technical PDF.

Your job is to identify every distinct figure, diagram, chart, flowchart, or image on the page
and produce a detailed, structured description of each one for use in a question-answering system.

For each figure you find, return a JSON object with these exact fields:
  - "label"       : the figure label exactly as written on the page (e.g. "Fig. 4", "Figure 2.1",
                    "Fig. 4. Integration of langchain framework in RAG for the proposed document QA system.").
                    If no label is visible, use "Unlabelled figure".
  - "description" : a thorough, self-contained description that includes ALL of the following:
                      * The type of diagram (flowchart, architecture diagram, bar chart, table, etc.)
                      * Every visible text label, annotation, step number, node name, and axis title
                      * Every arrow — what it connects and in which direction
                      * All decision branches and their outcomes (e.g. Yes/No, True/False, loop-back paths)
                      * Every numbered or lettered step and what it represents
                      * Any data values, legends, colour meanings, or percentage figures
                      * Any icons or symbols and what they represent (e.g. database icon, OpenAI logo)
                      * The overall process, workflow, or concept the figure communicates
                      * How the figure relates to the surrounding text on the page if visible
                    Write as if describing to someone who cannot see the image at all.
                    Be thorough — omitting a label or step means it becomes unsearchable.

Return ONLY a valid JSON array — no preamble, no explanation, no markdown fences.
If there are no figures on the page, return an empty array: []

Example format:
[
  {
    "label": "Fig. 4. Integration of langchain framework in RAG for the proposed document QA system.",
    "description": "An eight-step flowchart arranged left to right..."
  }
]
"""


async def _analyse_page_figures(
    png_bytes: bytes,
    page_num: int,
    file_name: str,
) -> List[dict]:
    """
    Send a full-page screenshot to GPT-4.1 mini and ask it to identify and
    describe every figure on that page.

    Returns a list of {label, description} dicts — one per figure found.
    Returns an empty list if the call fails or no figures are found.
    """
    try:
        from langchain_core.messages import HumanMessage, SystemMessage

        b64 = base64.standard_b64encode(png_bytes).decode("utf-8")

        messages = [
            SystemMessage(content=_FIGURE_ANALYSIS_SYSTEM_PROMPT),
            HumanMessage(
                content=[
                    {
                        "type": "image_url",
                        "image_url": {"url": f"data:image/png;base64,{b64}"},
                    },
                    {
                        "type": "text",
                        "text": (
                            f"This is page {page_num + 1} of the document '{file_name}'. "
                            "Please identify and describe every figure on this page."
                        ),
                    },
                ]
            ),
        ]

        response = await _vision_llm.ainvoke(messages)
        raw = response.content.strip()

        # Strip markdown fences if model adds them despite instructions
        if raw.startswith("```"):
            raw = "\n".join(
                line for line in raw.splitlines()
                if not line.strip().startswith("```")
            ).strip()

        # Attempt 1 — direct parse
        try:
            figures = json.loads(raw)
            if isinstance(figures, list):
                logger.info(
                    "Page %d of %s: %d figure(s) identified.",
                    page_num + 1, file_name, len(figures),
                )
                return figures
        except json.JSONDecodeError:
            pass

        # Attempt 2 — extract just the [...] array from the response in case
        # the model wrapped it in extra prose despite instructions
        start = raw.find("[")
        end   = raw.rfind("]")
        if start != -1 and end != -1 and end > start:
            try:
                figures = json.loads(raw[start : end + 1])
                if isinstance(figures, list):
                    logger.info(
                        "Page %d of %s: %d figure(s) identified (extracted array).",
                        page_num + 1, file_name, len(figures),
                    )
                    return figures
            except json.JSONDecodeError:
                pass

        # Attempt 3 — the model likely has unescaped quotes inside a description
        # string. Replace any " that follows a word character and precedes a word
        # character (i.e. mid-string) with a single quote as a best-effort fix,
        # then try parsing again.

        cleaned = re.sub(r'(?<=[a-zA-Z0-9,\.\!\?\:\;])"(?=[a-zA-Z0-9\s])', "'", raw)
        try:
            figures = json.loads(cleaned)
            if isinstance(figures, list):
                logger.info(
                    "Page %d of %s: %d figure(s) identified (after quote fix).",
                    page_num + 1, file_name, len(figures),
                )
                return figures
        except json.JSONDecodeError:
            pass

        # Final fallback — store the whole raw response as one unlabelled chunk
        # so the content is not lost even if we can't parse structure
        logger.warning(
            "Page %d: all JSON parse attempts failed — storing raw response as fallback.",
            page_num + 1,
        )
        return [{"label": "Unlabelled figure", "description": raw}]
    except Exception as exc:
        logger.warning(
            "Figure analysis failed for page %d of %s: %s",
            page_num + 1,
            file_name,
            exc,
        )
        return []

async def _ocr_page(png_bytes: bytes, page_num: int) -> Optional[str]:
    """
    Send a rasterized PDF page to GPT-4.1 mini and ask it to transcribe all text.
    Used only for scanned pages that have no extractable text layer.
    Returns the transcribed text, or None if the call fails.
    """
    try:
        b64 = base64.standard_b64encode(png_bytes).decode("utf-8")

        from langchain_core.messages import HumanMessage

        message = HumanMessage(
            content=[
                {
                    "type": "image_url",
                    "image_url": {"url": f"data:image/png;base64,{b64}"},
                },
                {
                    "type": "text",
                    "text": (
                        "This is a scanned document page. Transcribe ALL text you can read "
                        "exactly as it appears, preserving paragraphs, headings, bullet points, "
                        "and table structure where possible. "
                        "Do not summarise or paraphrase — output the raw text only."
                    ),
                },
            ]
        )

        response = await _vision_llm.ainvoke([message])
        text = response.content.strip()
        logger.info("OCR completed for scanned page %d (%d chars).", page_num, len(text))
        return text

    except Exception as exc:
        logger.warning("OCR failed for page %d: %s", page_num, exc)
        return None


def _image_bytes_to_png(raw_bytes: bytes) -> bytes:
    """Normalise any image format to PNG bytes for consistent base64 encoding."""
    with Image.open(io.BytesIO(raw_bytes)) as img:
        buf = io.BytesIO()
        img.convert("RGB").save(buf, format="PNG")
        return buf.getvalue()


# ---------------------------------------------------------------------------
# Table → Markdown
# ---------------------------------------------------------------------------


def _table_to_markdown(table_data: List[List[str]]) -> str:
    """Convert a 2-D list of cell strings to a GitHub-flavoured Markdown table."""
    if not table_data:
        return ""

    rows = [[str(cell or "").strip() for cell in row] for row in table_data]
    col_count = max(len(r) for r in rows)

    # Pad all rows to the same width
    rows = [r + [""] * (col_count - len(r)) for r in rows]

    header    = "| " + " | ".join(rows[0]) + " |"
    separator = "| " + " | ".join(["---"] * col_count) + " |"
    body      = "\n".join("| " + " | ".join(r) + " |" for r in rows[1:])

    return "\n".join([header, separator, body]) if body else "\n".join([header, separator])


# ---------------------------------------------------------------------------
# PDF extraction (text + images + tables)
# ---------------------------------------------------------------------------


def _extract_pdf_sync(file_path: str) -> dict:
    """
    Extract text, image pages, tables, and scanned pages from a PDF synchronously.
    Must be dispatched via run_in_executor.

    Image handling:
      - If a page has at least one qualifying image, the ENTIRE page is rasterized
        once at PAGE_RENDER_DPI. This captures all overlaid text labels, arrows,
        and vector elements that individual image extraction misses.
      - Multiple images on the same page → one screenshot, not multiple.
      - Pages with no text layer are treated as scanned → queued for OCR instead.

    Returns:
        {
          "text_docs":     [Document, ...],
          "image_pages":   [(page_num, png_bytes), ...],  # one entry per page
          "table_data":    [(page_num, [[str]]), ...],
          "scanned_pages": [(page_num, png_bytes), ...],
        }
    """
    text_docs    = []
    image_pages  = []
    table_data   = []
    scanned_pages = []

    ocr_scale  = OCR_RASTER_DPI / 72
    ocr_matrix = fitz.Matrix(ocr_scale, ocr_scale)

    page_scale  = PAGE_RENDER_DPI / 72
    page_matrix = fitz.Matrix(page_scale, page_scale)

    pdf_fitz = fitz.open(file_path)
    for page_num, page in enumerate(pdf_fitz):
        text = page.get_text("text").strip()

        if text:
            text_docs.append(
                Document(
                    page_content=text,
                    metadata={"page": page_num},
                )
            )
        else:
            # No text layer — rasterize whole page for OCR
            pix = page.get_pixmap(matrix=ocr_matrix)
            scanned_pages.append((page_num, pix.tobytes("png")))
            logger.info("Page %d has no text layer — queued for OCR.", page_num)
            continue   # scanned pages are not processed for images

        # Check if this page should be screenshotted:
        #   (a) has qualifying embedded image objects, OR
        #   (b) has significant vector drawing content (>20 paths) —
        #       catches flowcharts/diagrams drawn directly in the PDF
        #       that have no embedded image object at all (e.g. Fig. 1)
        qualifying = [
            img_info for img_info in page.get_images(full=True)
            if (
                pdf_fitz.extract_image(img_info[0]).get("width", 0)  >= MIN_IMAGE_WIDTH
                and pdf_fitz.extract_image(img_info[0]).get("height", 0) >= MIN_IMAGE_HEIGHT
            )
        ]

        drawing_count = len(page.get_drawings())
        has_images    = len(qualifying) > 0
        has_drawings  = drawing_count > 20   # threshold avoids plain text pages with borders

        if has_images or has_drawings:
            pix = page.get_pixmap(matrix=page_matrix)
            image_pages.append((page_num, pix.tobytes("png")))
            logger.info(
                "Page %d: screenshotted [embedded=%d, drawing_paths=%d].",
                page_num + 1, len(qualifying), drawing_count,
            )

    pdf_fitz.close()

    # Tables via pdfplumber
    with pdfplumber.open(file_path) as pdf:
        for page_num, page in enumerate(pdf.pages):
            for table in page.extract_tables():
                if table:
                    table_data.append((page_num, table))

    logger.info(
        "PDF extraction complete: %d text page(s), %d image page(s), "
        "%d table(s), %d scanned page(s).",
        len(text_docs), len(image_pages), len(table_data), len(scanned_pages),
    )

    return {
        "text_docs":     text_docs,
        "image_pages":   image_pages,
        "table_data":    table_data,
        "scanned_pages": scanned_pages,
    }


# ---------------------------------------------------------------------------
# DOCX extraction (text + images + tables)
# ---------------------------------------------------------------------------


def _extract_docx_sync(file_path: str) -> dict:
    """
    Extract text, images, and tables from a DOCX synchronously.
    Must be dispatched via run_in_executor.
    """
    text_docs  = []
    image_data = []
    table_data = []

    doc = DocxDocument(file_path)

    # --- Text (paragraph-level) ---------------------------------------------
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    if paragraphs:
        text_docs.append(
            Document(
                page_content="\n".join(paragraphs),
                metadata={"page": 0},
            )
        )

    # --- Tables -------------------------------------------------------------
    for table in doc.tables:
        rows = [[cell.text for cell in row.cells] for row in table.rows]
        if rows:
            table_data.append((0, rows))

    # --- Images (embedded in document relationships) -----------------------
    # DOCX has no page rendering concept so we extract raw image bytes directly
    docx_images = []
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            try:
                raw = rel.target_part.blob
                png = _image_bytes_to_png(raw)
                docx_images.append((0, png))
            except Exception as exc:
                logger.warning("Skipping DOCX image (extraction error): %s", exc)

    return {
        "text_docs":     text_docs,
        "image_pages":   docx_images,   # treated as "pages" — one chunk per image
        "table_data":    table_data,
        "scanned_pages": [],
    }


# ---------------------------------------------------------------------------
# TXT extraction (text only — images/tables not applicable)
# ---------------------------------------------------------------------------


def _extract_txt_sync(file_path: str) -> dict:
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        content = f.read()
    return {
        "text_docs":     [Document(page_content=content, metadata={"page": 0})],
        "image_pages":   [],
        "table_data":    [],
        "scanned_pages": [],
    }


# ---------------------------------------------------------------------------
# Dispatch
# ---------------------------------------------------------------------------


_EXTRACTORS = {
    "pdf":  _extract_pdf_sync,
    "docx": _extract_docx_sync,
    "txt":  _extract_txt_sync,
}


async def _extract_multimodal(file_path: str, file_type: str) -> dict:
    """Non-blocking dispatch to the correct sync extractor."""
    if file_type not in _EXTRACTORS:
        raise ValueError(f"Unsupported file type: {file_type!r}. Supported: {SUPPORTED_TYPES}")
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(None, _EXTRACTORS[file_type], file_path)


# ---------------------------------------------------------------------------
# Build LangChain Documents from extracted multimodal content
# ---------------------------------------------------------------------------


async def _build_all_chunks(extracted: dict, file_name: str) -> List[Document]:
    """
    Convert raw extracted content into LangChain Documents ready for embedding.

    Text          → split with RecursiveCharacterTextSplitter (modality="text")
    Image pages   → full-page screenshot analysed by GPT-4.1 mini; each figure
                    identified becomes its own chunk (modality="image")
    Tables        → convert to Markdown, stored as a chunk (modality="table")
    Scanned pages → transcribed via GPT-4.1 mini OCR, split (modality="ocr")
    """
    chunks: List[Document] = []

    # 1. Text chunks ---------------------------------------------------------
    text_chunks = _splitter.split_documents(extracted["text_docs"])
    for chunk in text_chunks:
        chunk.metadata.setdefault("modality", "text")
    chunks.extend(text_chunks)
    logger.info("Text: %d chunk(s) after splitting.", len(text_chunks))

    # 2. Image chunks — one LLM call per page, returns per-figure JSON -------
    image_tasks = [
        _analyse_page_figures(png_bytes, page_num, file_name)
        for page_num, png_bytes in extracted["image_pages"]
    ]
    page_results = await asyncio.gather(*image_tasks)

    image_chunk_count = 0
    for (page_num, _), figures in zip(extracted["image_pages"], page_results):
        for fig in figures:
            label       = fig.get("label", "Unlabelled figure")
            description = fig.get("description", "").strip()
            if not description:
                continue
            chunks.append(
                Document(
                    page_content=(
                        f"[Figure on page {page_num + 1} — {label}]: {description}"
                    ),
                    metadata={"modality": "image", "page": page_num},
                )
            )
            image_chunk_count += 1

    logger.info(
        "Images: %d page(s) analysed, %d figure chunk(s) produced.",
        len(extracted["image_pages"]),
        image_chunk_count,
    )

    # 3. Table chunks --------------------------------------------------------
    for page_num, table_rows in extracted["table_data"]:
        md = _table_to_markdown(table_rows)
        if not md.strip():
            continue
        chunks.append(
            Document(
                page_content=f"[Table on page {page_num + 1}]:\n{md}",
                metadata={"modality": "table", "page": page_num},
            )
        )
    logger.info("Tables: %d chunk(s) added.", len(extracted["table_data"]))

    # 4. Scanned page chunks (OCR) -------------------------------------------
    ocr_tasks = [
        _ocr_page(png_bytes, page_num)
        for page_num, png_bytes in extracted["scanned_pages"]
    ]
    ocr_results = await asyncio.gather(*ocr_tasks)

    for (page_num, _), ocr_text in zip(extracted["scanned_pages"], ocr_results):
        if not ocr_text:
            continue
        ocr_doc = Document(
            page_content=ocr_text,
            metadata={"page": page_num},
        )
        ocr_chunks = _splitter.split_documents([ocr_doc])
        for chunk in ocr_chunks:
            chunk.metadata["modality"] = "ocr"
        chunks.extend(ocr_chunks)

    logger.info(
        "Scanned pages: %d queued, %d transcribed successfully.",
        len(extracted["scanned_pages"]),
        sum(1 for r in ocr_results if r),
    )

    return chunks


# ---------------------------------------------------------------------------
# Embed + upsert (unchanged from original, works for all modalities)
# ---------------------------------------------------------------------------


@_with_retry()
async def _embed_and_upsert(chunks: List[Document], namespace: str) -> None:
    total = len(chunks)
    loop  = asyncio.get_event_loop()

    logger.info("Upserting %d chunk(s)...", total)

    for start in range(0, total, EMBED_BATCH_SIZE):
        batch = chunks[start : start + EMBED_BATCH_SIZE]

        await loop.run_in_executor(
            None,
            lambda b=batch: vector_store.add_documents(b, namespace=namespace),
        )

        logger.info(
            "  Upserted %d-%d / %d",
            start + 1, min(start + EMBED_BATCH_SIZE, total), total,
        )


# ---------------------------------------------------------------------------
# Main entry point
# ---------------------------------------------------------------------------


async def process_document(document_id: str) -> None:
    """
    Full multimodal pipeline:
      Load → Extract (text + images + tables) → Build chunks →
      Filter → Inject metadata → Embed (batched) → Upsert → Update status
    """
    document   = None
    start_time = time.monotonic()

    try:
        # 0. Fetch record ----------------------------------------------------
        document = await DocumentModel.get(document_id)

        if not document:
            logger.warning("Unknown document_id=%s — skipping.", document_id)
            return

        namespace = str(document.user_id)
        logger.info("Starting: document_id=%s  type=%s", document_id, document.file_type)

        document.status     = "processing"
        document.updated_at = _utcnow()
        await document.save()

        # 1. File size guard -------------------------------------------------
        import os
        file_size = os.path.getsize(document.file_path)
        if file_size > MAX_FILE_SIZE_BYTES:
            raise ValueError(
                f"File size {file_size / 1024 / 1024:.1f} MB exceeds limit of {MAX_FILE_SIZE_MB} MB."
            )

        # 2. Extract all modalities (non-blocking) ---------------------------
        extracted = await _extract_multimodal(document.file_path, document.file_type)
        logger.info(
            "Extracted: %d text page(s), %d image page(s), %d table(s), %d scanned page(s).",
            len(extracted["text_docs"]),
            len(extracted["image_pages"]),
            len(extracted["table_data"]),
            len(extracted["scanned_pages"]),
        )

        # 3. Build chunks across all modalities (image captioning is async) --
        file_name = document.file_path.split("/")[-1]
        chunks = await _build_all_chunks(extracted, file_name)
        logger.info("Total chunks before filtering: %d", len(chunks))

        # 4. Chunk count guard -----------------------------------------------
        if len(chunks) > MAX_CHUNKS_GUARD:
            raise ValueError(
                f"Document yielded {len(chunks)} chunks — exceeds guard of {MAX_CHUNKS_GUARD}."
            )

        # 5. Filter empty chunks ---------------------------------------------
        chunks = _filter_chunks(chunks)
        if not chunks:
            raise ValueError("All chunks were empty after filtering.")

        # 6. Inject metadata (document_id → chat_id → user_id → modality → page)
        chunks = _inject_metadata(chunks, document)

        # 7. Embed + upsert --------------------------------------------------
        await _embed_and_upsert(chunks, namespace)

        # 8. Persist success -------------------------------------------------
        document.status      = "processed"
        document.chunk_count = len(chunks)
        document.updated_at  = _utcnow()
        await document.save()

        logger.info(
            "Done: document_id=%s  chunks=%d  elapsed=%.2fs",
            document_id, len(chunks), time.monotonic() - start_time,
        )

    except Exception as exc:
        logger.exception(
            "Failed: document_id=%s  elapsed=%.2fs  error=%s",
            document_id, time.monotonic() - start_time, exc,
        )
        if document:
            document.status        = "failed"
            document.error_message = str(exc)
            document.updated_at    = _utcnow()
            await document.save()

